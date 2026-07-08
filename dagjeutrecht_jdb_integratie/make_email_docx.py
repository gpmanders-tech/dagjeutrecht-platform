"""Word-bestand met twee uitnodigingsmails (JdB + Schuttevaer)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm

OUT = (
    r"C:\Users\gpman\OneDrive - Manders Travel & Events\Kantoor Manders - ETS"
    r"\10. Claudine\Mails_partneruitnodiging_dagjeutrecht.docx"
)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2); s.right_margin = Cm(2)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def h(text, level=1):
    head = doc.add_heading(text, level=level)
    for r in head.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def p(text="", bold=False, italic=False):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    return para


def meta(label, value=""):
    para = doc.add_paragraph()
    r = para.add_run(f"{label} "); r.bold = True
    para.add_run(value)


# ============== TITEL ==============
title = doc.add_heading("Uitnodigingsmails partners — dagjeutrecht.nl koppeling", level=0)
for r in title.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

p(
    "Twee korte uitnodigingsmails om met onze huidige partners het gesprek te openen over "
    "een technische koppeling met dagjeutrecht.nl. Doel: bestaande samenwerking "
    "professionaliseren, niet meteen alle techniek bespreken.",
    italic=True,
)
doc.add_paragraph()

# ============== MAIL 1: JDB ==============
h("Mail 1 — Jeu de Boules Utrecht", level=1)

meta("Onderwerp:", "Onze samenwerking naar een volgend niveau – voorstel voor kort overleg")
meta("Aan:", "[contactpersoon JdB Utrecht]")
meta("Van:", "Ger Manders – Manders Travel / dagjeutrecht.nl")
doc.add_paragraph()

p("Beste [naam],")
doc.add_paragraph()

p(
    "We werken inmiddels prettig samen rond de jeu-de-boules-arrangementen die we via "
    "dagjeutrecht.nl en onze andere kanalen bij groepen onder de aandacht brengen. "
    "Ik merk dat de vraag groeit – vooral combinaties met andere actieve onderdelen "
    "(suppen, kanoën, borrelarrangement) doen het goed."
)
doc.add_paragraph()

p(
    "Graag wil ik met jullie verkennen of we deze samenwerking een stap verder kunnen "
    "brengen door dagjeutrecht.nl en jullie boekingssysteem (Smart Event Manager) "
    "technisch te koppelen. Concreet stel ik me voor:"
)
for line in [
    "Groepen boeken een compleet arrangement (bv. sup + JdB + borrel) in één keer op dagjeutrecht.nl",
    "Manders factureert de klant; wij rekenen achteraf met jullie af",
    "Jullie zien jullie deel van de boeking direct in SEM verschijnen – geen losse mails en handmatige administratie meer",
    "Meer volume via dagjeutrecht.nl, met behoud van jullie eigen agenda en regie",
]:
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(line)

doc.add_paragraph()
p(
    "Voor jullie betekent het: minder afstemming per boeking, professionelere uitstraling "
    "richting groepen, en een vast verkoopkanaal dat actief promoot. Voor ons: één "
    "naadloze ervaring voor de klant in plaats van losse stukjes."
)
doc.add_paragraph()

p(
    "Zouden we daar een kort overleg over kunnen plannen – zo'n drie kwartier? "
    "Ik kom graag bij jullie langs op de baan. Een paar mogelijke momenten: "
    "[datum 1], [datum 2] of [datum 3]. Werkt geen daarvan, dan stuur ik graag een paar "
    "alternatieven."
)
doc.add_paragraph()

p("Hartelijke groet,")
p("Ger Manders")
p("Manders Travel & Events | dagjeutrecht.nl")
p("06 - [nummer] | info@manatooevents.nl")

doc.add_page_break()

# ============== MAIL 2: SCHUTTEVAER ==============
h("Mail 2 — Schuttevaer Rondvaartboten", level=1)

meta("Onderwerp:", "Samen meer groepen aan boord – voorstel voor kort overleg")
meta("Aan:", "[contactpersoon Schuttevaer]")
meta("Van:", "Ger Manders – Manders Travel / dagjeutrecht.nl")
doc.add_paragraph()

p("Beste [naam],")
doc.add_paragraph()

p(
    "Onze rondvaarten via Schuttevaer zijn een vaste waarde geworden in het aanbod dat "
    "we vanuit dagjeutrecht.nl naar groepen brengen – bedrijfsuitjes, vrijgezellenfeesten, "
    "schoolgroepen. We merken dat veel klanten een rondvaart het liefst combineren met "
    "iets actiefs vooraf (suppen, kanoën) of een borrel achteraf."
)
doc.add_paragraph()

p(
    "Graag wil ik met jullie verkennen of we de samenwerking technisch wat steviger "
    "kunnen verankeren via een directe koppeling tussen dagjeutrecht.nl en jullie "
    "FareHarbor-omgeving. Wat ik voor ogen heb:"
)
for line in [
    "Groepen boeken in één flow een combi-arrangement op dagjeutrecht.nl (vaart + extra's)",
    "Beschikbaarheid van jullie boten is realtime zichtbaar op onze site",
    "Boekingen verschijnen direct in jullie FareHarbor – zonder dat iemand iets hoeft over te tikken",
    "Manders factureert de klant en regelt de afrekening met jullie",
]:
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(line)

doc.add_paragraph()
p(
    "Voor jullie scheelt het administratie en mailverkeer per boeking, en levert het een "
    "vast kanaal op dat actief groepen werft. FareHarbor leent zich technisch goed voor "
    "dit type koppeling, dus de stap is naar verwachting overzichtelijk."
)
doc.add_paragraph()

p(
    "Kunnen we hier op korte termijn een halfuurtje voor inplannen? Ik kan ook even "
    "langskomen aan de Oudegracht. Een paar mogelijke momenten: [datum 1], [datum 2] of "
    "[datum 3]. Als niets daarvan past, hoor ik graag wat wél schikt."
)
doc.add_paragraph()

p("Hartelijke groet,")
p("Ger Manders")
p("Manders Travel & Events | dagjeutrecht.nl")
p("06 - [nummer] | info@manatooevents.nl")

doc.add_page_break()

# ============== TIPS ==============
h("Tips bij verzenden", level=1)
for line in [
    "Pas [naam] en de drie data aan voor je verstuurt.",
    "Stuur niet beide mails op dezelfde dag — laat 1–2 weken tussenruimte tussen de gesprekken om geleerde lessen mee te nemen.",
    "Begin met Schuttevaer (FareHarbor) – technisch het laagdrempelige pad, en goede pilot voor het patroon richting JdB.",
    "Stuur de intake-checklist NIET mee als bijlage. Die is voor het overleg zelf; vooraf werkt het te zwaar.",
    "Bij positief signaal: plan meteen het overleg, gebruik de intake-checklist als agenda.",
]:
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(line)

doc.save(OUT)
print(f"Geschreven: {OUT}")
