"""Bouw één Word-bestand met beide intake-checklists (JdB + Schuttevaer)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUT_DIR = r"C:\Users\gpman\OneDrive - Manders Travel & Events\Kantoor Manders - ETS\10. Claudine"
OUT_FILE = rf"{OUT_DIR}\Intake_checklists_dagjeutrecht_partners.docx"

doc = Document()

# Stel marges in
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(text, checkbox=True):
    p = doc.add_paragraph(style="List Bullet")
    prefix = "☐ " if checkbox else ""
    p.add_run(prefix + text)
    return p


def add_sub(text):
    p = doc.add_paragraph(style="List Bullet 2")
    p.add_run("☐ " + text)
    return p


def add_meta_line(label):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True


def add_table_2col(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, cell in enumerate(row):
            t.rows[ri].cells[ci].text = cell


# ============== TITEL ==============
title = doc.add_heading("Intake-checklists partnerkoppelingen dagjeutrecht.nl", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

p = doc.add_paragraph()
p.add_run(
    "Twee voorbereidende checklists voor partnergesprekken om te komen tot een "
    "boekingsintegratie op dagjeutrecht.nl. Klant boekt op 1 site, betaalt 1 factuur "
    "(WeFact) aan Manders, partners worden achteraf afgerekend."
).italic = True

doc.add_paragraph()

# ============== DEEL 1: JDB ==============
add_heading("Deel 1 — Jeu de Boules Utrecht (Smart Event Manager)", level=1)

add_para(
    "Doel: dagjeutrecht.nl verkoopt combi-arrangementen (sup/kano + JdB met borrel). "
    "JdB werkt met Smart Event Manager (SEM) — koppeling via SEM Connector/API."
)
add_meta_line("Te bespreken met: Jeu de Boules Utrecht (eigenaar + technisch contact SEM)")
add_meta_line("Datum overleg:")
add_meta_line("Aanwezig:")

# 1. Commercieel
add_heading("1. Commercieel & samenwerking", level=2)
for item in [
    "Akkoord op reseller-constructie (dagjeutrecht verkoopt onder eigen merk)?",
    "Welke producten mag dagjeutrecht aanbieden? (banen los / met borrel / hapjes / volledige arrangementen)",
    "Inkoopprijzen voor dagjeutrecht per product (incl./excl. btw)",
    "Vaste arrangementsprijs of stapelen we per onderdeel?",
    "Minimum / maximum groepsgrootte per product",
    "Tijdsblokken (vaste blokken of flexibel)",
    "Exclusiviteit — mag dagjeutrecht JdB verkopen of werken er meer partners?",
    "Geldigheidsduur prijsafspraak (1 jaar? seizoen?)",
]:
    add_bullet(item)

# 2. Operationeel
add_heading("2. Operationeel", level=2)
for item in [
    "Openingstijden + uitzonderingen (feestdagen, sluitingsperiodes)",
    "Hoe ver vooruit kan geboekt worden?",
    "Cut-off: tot hoe laat vóór aanvang nog boekbaar?",
    "Wie ontvangt de boekingsbevestiging bij JdB? (mail + telefoon)",
    "Wie is dag-contact bij wijziging of no-show?",
    "Hoe gaat JdB om met slecht weer (sup/kano gaat dan misschien niet door)?",
    "Communicatie naar de gast — stuurt dagjeutrecht alles of JdB ook?",
]:
    add_bullet(item)

# 3. Financieel
add_heading("3. Financieel & afrekening", level=2)
for item in [
    "Hoe wil JdB afgerekend worden? (per boeking, wekelijks, maandelijks)",
    "Betaaltermijn factuur JdB → Manders",
    "BTW: hoog/laag tarief per onderdeel (banen vs. horeca)",
    "Annuleringsvoorwaarden JdB → Manders (staffel %)",
    "Wat bij no-show van gast? (volledige doorbelasting?)",
    "Aanbetaling/borg vereist door JdB richting Manders?",
    "Bar-afrekening (consumpties boven arrangement) — direct of via Manders?",
]:
    add_bullet(item)

# 4. SEM technisch
add_heading("4. SEM-toegang (technisch — kern van de bouw)", level=2)
for item in [
    "Welke editie/versie van Smart Event Manager gebruikt JdB?",
    "Heeft JdB de Connector-licentie actief? (zo niet: kosten + activatietijd opvragen)",
    "API-credentials beschikbaar? (key / secret / endpoint URL)",
    "Is er een test/sandbox-omgeving of alleen live?",
    "Documentatie beschikbaar (PDF, link, of via SEM-support)?",
]:
    add_bullet(item)

add_para("Welke endpoints zijn beschikbaar:", bold=True)
for item in [
    "beschikbaarheid opvragen (banen + arrangement op datum/tijd)",
    "reservering aanmaken",
    "reservering wijzigen",
    "reservering annuleren",
    "producten/prijzen ophalen",
]:
    add_sub(item)

for item in [
    "Ondersteunt SEM webhooks (notificatie bij wijziging vanuit JdB-kant)?",
    "Lopen er al andere koppelingen op deze SEM-installatie?",
    "Technisch contact bij SEM-leverancier (naam + mail)",
    "Technisch contact bij JdB (wie kent het systeem) — naam + mail",
]:
    add_bullet(item)

# 5. Juridisch
add_heading("5. Juridisch & risico", level=2)
for item in [
    "Samenwerkings-/reseller-overeenkomst vastleggen (kort document)",
    "Aansprakelijkheid bij ongeval/schade tijdens JdB-activiteit",
    "JdB verzekering WA / ongevallen — bewijs opvragen",
    "AVG: welke klantgegevens deelt Manders met JdB? Verwerkersovereenkomst nodig?",
    "Pakketreizenrichtlijn / garantiefonds — toetsen of combi onder pakketreis valt",
    "Wie staat op de factuur naar eindklant? (Manders — niet JdB)",
]:
    add_bullet(item)

# 6. Marketing
add_heading("6. Marketing & content", level=2)
for item in [
    "Mag dagjeutrecht foto's/teksten van JdB-locatie gebruiken?",
    "Hoe wordt JdB in de productbeschrijving genoemd (wel/niet bij naam)?",
    "Reviews/feedback — wie verzamelt en deelt?",
]:
    add_bullet(item)

add_heading("Notities tijdens overleg JdB", level=2)
for _ in range(6):
    doc.add_paragraph()

doc.add_page_break()

# ============== DEEL 2: SCHUTTEVAER ==============
add_heading("Deel 2 — Schuttevaer Rondvaartboten (FareHarbor)", level=1)

add_para(
    "Doel: dagjeutrecht.nl biedt rondvaart van Schuttevaer aan als los product óf in combi. "
    "Schuttevaer werkt met FareHarbor — bekend uit dagjesuppen-project, technisch het "
    "eenvoudigste pad."
)
add_meta_line("Te bespreken met: Schuttevaer (eigenaar + operationeel contact)")
add_meta_line("Datum overleg:")
add_meta_line("Aanwezig:")

# 1. Commercieel
add_heading("1. Commercieel & samenwerking", level=2)
for item in [
    "Akkoord op reseller-constructie?",
    "Welke vaarproducten mag dagjeutrecht aanbieden? (open rondvaart / private charter / catering / themavaart)",
    "Inkoopprijzen per producttype (incl./excl. btw)",
    "Combi-opties: borrel/hapjes/lunch aan boord — apart of in pakket?",
    "Capaciteit per boot (min/max personen) en aantal boten",
    "Vaartijden / duur (vaste blokken? maatwerk?)",
    "Vaargebied + opstap-/uitstapplek (vast of meerdere?)",
    "Exclusiviteit",
    "Geldigheidsduur prijsafspraak",
]:
    add_bullet(item)

# 2. Operationeel
add_heading("2. Operationeel", level=2)
for item in [
    "Vaartseizoen (jaarrond? alleen apr–okt?)",
    "Hoe ver vooruit boekbaar?",
    "Cut-off voor laatste boeking vóór vertrek",
    "Slechtweer-/laagwater-beleid — wanneer afgelast? Vergoeding of verzetten?",
    "Wie ontvangt boekingsbevestiging bij Schuttevaer?",
    "Wie is dag-contact bij wijziging / no-show / opstap-probleem?",
    "Communicatie naar gast (opstap, parkeren, weer) — wie stuurt?",
    "Schippers/bemanning — geregeld door Schuttevaer, of moet dagjeutrecht iets weten?",
]:
    add_bullet(item)

# 3. Financieel
add_heading("3. Financieel & afrekening", level=2)
for item in [
    "Hoe wil Schuttevaer afgerekend worden? (per boeking, wekelijks, maandelijks)",
    "Betaaltermijn factuur Schuttevaer → Manders",
    "BTW: 9% (personenvervoer) of 21% (charter/horeca)? Per onderdeel uitsplitsen?",
    "Annuleringsvoorwaarden Schuttevaer → Manders (staffel %)",
    "Wat bij no-show?",
    "Wat bij afgelasten door Schuttevaer? Volledige restitutie aan Manders?",
    "Aanbetaling/borg richting Schuttevaer?",
    "Consumpties aan boord boven arrangement — direct door gast of via Manders?",
]:
    add_bullet(item)

# 4. FareHarbor technisch
add_heading("4. FareHarbor-toegang (technisch)", level=2)
for item in [
    "FareHarbor account-naam (shortname) van Schuttevaer",
    "API-key beschikbaar? (Schuttevaer maakt aan in FH-dashboard)",
]:
    add_bullet(item)

add_para(
    "Heeft Schuttevaer een affiliate/reseller-programma in FareHarbor? "
    "Dit is technisch de cleanste route:",
    bold=True,
)
for item in [
    "Manders krijgt eigen affiliate-account",
    "Boekingen lopen direct via FH API met affiliate-tracking",
    "Commissie automatisch verrekend",
]:
    add_sub(item)

for item in [
    "Welke items/availabilities zijn via API beschikbaar? Lijst per producttype.",
    "Ondersteunt FH-setup custom fields (allergieën, opmerkingen)?",
    "Welke velden zijn verplicht bij booking-create?",
    "Webhooks geconfigureerd (annulering vanuit Schuttevaer)?",
    "Lopen er al andere kanalen op deze FH-installatie?",
]:
    add_bullet(item)

p = doc.add_paragraph()
r = p.add_run(
    "Tip: FareHarbor heeft uitstekende publieke docs (fareharbor.com/help/external-api). "
    "Veel vragen kunnen we technisch zelf beantwoorden zodra we shortname + API-key hebben."
)
r.italic = True

# 5. Juridisch
add_heading("5. Juridisch & risico", level=2)
for item in [
    "Samenwerkings-/reseller-overeenkomst vastleggen",
    "Aansprakelijkheid tijdens vaart — wie is verantwoordelijk?",
    "Schuttevaer verzekering (P&I / passagiers) — bewijs opvragen",
    "Veiligheidscertificaten boten + schippersvaarbewijs — bewijs/kopie",
    "AVG: welke klantgegevens deelt Manders met Schuttevaer?",
    "Pakketreizenrichtlijn / garantiefonds — toetsen (vooral bij multi-day combi's)",
    "Factuur naar eindklant op naam van Manders",
]:
    add_bullet(item)

# 6. Marketing
add_heading("6. Marketing & content", level=2)
for item in [
    "Foto's/video's van boten + vaargebied — mag dagjeutrecht gebruiken?",
    "Schuttevaer wel/niet bij naam noemen in productbeschrijving?",
    "Reviews — gedeeld of apart?",
]:
    add_bullet(item)

add_heading("Notities tijdens overleg Schuttevaer", level=2)
for _ in range(6):
    doc.add_paragraph()

doc.add_page_break()

# ============== VERGELIJKING ==============
add_heading("Vergelijking & volgorde-advies", level=1)

add_table_2col(
    ["Aspect", "JdB (SEM)", "Schuttevaer (FareHarbor)"],
    [],
)
# add_table_2col only takes 2 col header, replace with proper 3-col table
# Remove the empty one we just added (last table)
# (python-docx heeft geen 'remove' op tabel niveau via API; we doen het opnieuw met directe call)

# Vorige (lege) tabel verwijderen via XML
last_tbl = doc.tables[-1]._element
last_tbl.getparent().remove(last_tbl)

# Correcte 3-koloms tabel
t = doc.add_table(rows=5, cols=3)
t.style = "Light Grid Accent 1"
rows_data = [
    ["Aspect", "JdB (SEM)", "Schuttevaer (FareHarbor)"],
    ["API-volwassenheid", "Moeilijker (Connector, weinig publieke docs)",
     "Goed (publieke docs, bekend uit dagjesuppen)"],
    ["Affiliate-model", "Onwaarschijnlijk", "Mogelijk → mogelijk cleanere oplossing"],
    ["Annulering", "Reseller-deal", "Idem, maar FH heeft annuleringsflow ingebouwd"],
    ["Voorbeeld-implementatie", "Nieuw bouwen", "Hergebruik dagjesuppen.nl-patroon"],
]
for ri, row in enumerate(rows_data):
    for ci, val in enumerate(row):
        c = t.rows[ri].cells[ci]
        c.text = val
        if ri == 0:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True

add_para("")
add_para("Volgorde-advies:", bold=True)
add_para(
    "Start technisch met Schuttevaer (FareHarbor) — hergebruik van het dagjesuppen-werk "
    "maakt MVP sneller live. JdB-koppeling daarna, want de SEM Connector-API moet eerst "
    "verkend worden (en is afhankelijk van JdB's licentie/medewerking)."
)

add_heading("Vervolgstappen na intakegesprekken", level=2)
for item in [
    "Antwoorden verwerken in beslissingsdocument (1 A4 per partner)",
    "Technische credentials opvragen (FH-shortname + API-key voor Schuttevaer; SEM-Connector voor JdB)",
    "Eerste API-test uitvoeren (Schuttevaer eerst)",
    "Bouw-prompt schrijven voor Claude met verzamelde gegevens",
    "Korte samenwerkingsovereenkomst opstellen per partner",
]:
    add_bullet(item, checkbox=True)

# Opslaan
doc.save(OUT_FILE)
print(f"Geschreven: {OUT_FILE}")
